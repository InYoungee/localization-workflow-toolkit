import os
import glob
import json
import xml.etree.ElementTree as ET
from docx import Document
import PyPDF2
import pandas as pd
from datetime import datetime

def extract_text_from_json(filepath):
	"""Extract text from JSON file"""
	def extract_strings(obj):
		text = []
		if isinstance(obj, dict):
			for value in obj.values():
				text.extend(extract_strings(value))
		elif isinstance(obj, list):
			for item in obj:
				text.extend(extract_strings(item))
		elif isinstance(obj, str):
			text.append(obj)
		return text

	try:
		with open(filepath, 'r', encoding='utf-8') as file:
			data = json.load(file)
		all_text = extract_strings(data)
		return ' '.join(all_text)
	except Exception as e:
		print(f" X Error reading JSON: {e}")
		return ""

def extract_text_from_xml(filepath):
	"""Extract text from XML/XLF file
	For XLF files: extract only source text (not target)
	for other XML: extrac all text
	"""
	try:
		tree = ET.parse(filepath)
		root = tree.getroot()

		all_text = []

		# Check if this is an XLIFF file
		is_xliff = False
		if 'xliff' in root.tag.lower() or any('xliff' in elem.tag.lower() for elem in root.iter()):
			is_xliff = True
		if is_xliff:
			namespaces = {'xliff': 'urn:oasis:names:tc:xliff:document:1.2'}

			sources = root.findall('.//xliff:source', namespaces)

			if not sources:
				sources = root.findall('//{*}source')
			if not sources:
				for elem in root.iter():
					if 'source' in elem.tag.lower():
						if elem.text and elem.text.strip():
							all_text.append(elem.text.strip())
			else:
				for source in sources:
					if source.text and source.text.strip():
						all_text.append(source.text.strip())
		else:
			for elem in root.iter():
				if elem.text and elem.text.strip():
					all_text.append(elem.text.strip())
				if elem.tail and elem.tail.strip():
					all_text.append(elem.tail.strip())
		return ' '.join(all_text)

	except Exception as e:
		print(f" X Error reading XML: {e}")
		return ""


def extract_text_from_docx(filepath):
	"""Extract text from Word doc"""
	try:
		doc = Document(filepath)
		text = []
		#Extract text from paragraph
		for paragraph in doc.paragraphs:
			if paragraph.text.strip():
				text.append(paragraph.text)
		#Extract text from tables
		for table in doc.tables:
			for row in table.rows:
				for cell in row.cells:
					if cell.text.strip():
						text.append(cell.text)

		return ' '.join(text)
	except Exception as e:
		print(f" X Error reading DOCX: {e}")
		return ""

def extract_text_from_pdf(filepath):
	"""Extract text from PDF"""
	try:
		text = []
		with open(filepath, 'rb') as file:
			pdf_reader = PyPDF2.PdfReader(file)

			#Extract text from each page
			for page_num in range(len(pdf_reader.pages)):
				page = pdf_reader.pages[page_num]
				page_text = page.extract_text()
				if page_text.strip():
					text.append(page_text)

		return ' '.join(text)
	except Exception as e:
		print(f" X Error reading PDF: {e}")
		return ""

def count_words_in_file(filepath):
	filename = os.path.basename(filepath)
	_, ext = os.path.splitext(filename)
	ext = ext.lower()

	print(f" Processing: {filename}")
	#Route to appropriate extractor based on file type
	if ext == '.json':
		text = extract_text_from_json(filepath)
		file_type = 'JSON'
	elif ext in ['.xml', '.xlf']:
		text = extract_text_from_xml(filepath)
		file_type = 'XML/XLF'
	elif ext == '.docx':
		text = extract_text_from_docx(filepath)
		file_type = 'DOCX'
	elif ext == '.pdf':
		text = extract_text_from_pdf(filepath)
		file_type = 'PDF'
	else:
		print(f" ⚠️ Unsupported file type: {ext}")
		return None

	# Count words
	if text:
		words = len(text.split())
		print(f" ✓ {words:,} words")
		return {
			'filename': filename,
			'file_type': file_type,
			'words': words
		}
	else:
		print(f" ⚠️ No text extracted")
		return None

def analyze_folder(folder_path, file_patterns=None):
	"""Analyze all supported files in a folder
	Args:
		folder_path: Path to folder containing files
		file_patterns: List of patterns like ['*.json', '*.xml']
		If None, searches for all supported types
	"""

	# Default patterns for all supported types
	if file_patterns is None:
		file_patterns = ['*.json', '*.xml','*.xlf', '*.docx', '*.pdf']

	# Find all matching files
	all_files = []
	for pattern in file_patterns:
		search_path = os.path.join(folder_path, pattern)
		all_files.extend(glob.glob(search_path))

	if not all_files:
		print(f"\nX No supported files found in {folder_path}")
		print(f" Looking for: {', '.join(file_patterns)}")
		return None

	print(f"\n{'='*70}")
	print(f"ANALYZING FOLDER: {folder_path}")
	print(f"\n{'=' * 70}")
	print(f"Found {len(all_files)} file(s_ to analyze\n")

	results = []

	# Process each file
	for filepath in all_files:
		result = count_words_in_file(filepath)
		if result:
			results.append(result)
		print()

	if not results:
		print("X No files processed successfully")
		return None
	return results

def display_summary(results, cost_per_word=0.15):
	if not results:
		return
	print("="*70)
	print("SUMMARY REPORT")
	print("="*70)
	print(f"{'File Name':<35} {'Type':<12} {'Words':>10}")
	print("="*70)

	total_words = 0

	# Sort by file type, then by words (descending)
	results.sort(key=lambda x: (x['file_type'], -x['words']))

	for result in results:
		print(f"{result['filename']:<35} {result['file_type']:<12} {result['words']:>10,}")
		total_words += result['words']

	print("=" * 70)
	print(f"{'TOTAL':<35} {'':<12} {total_words:>10,}")
	print("=" * 70)

	estimated_cost = total_words * cost_per_word
	estimated_hours = total_words / 250 # 250 words per hour

	print(f"\nTotal files processed: {len(results)}")
	print(f"Total words: {total_words:,}")
	print(f"Estimated translation cost: ${estimated_cost:,.2f} (at ${cost_per_word}/word)")
	print(f"Estimated time: {estimated_hours:.1f} hours (at 250 words/hour)")
	print("=" * 70)

def export_to_excel(results, cost_per_word=0.15):
	if not results:
		return

	data = []
	for result in results:
		data.append({
			'File Name': result['filename'],
			'Type': result['file_type'],
			'Words': result['words'],
			'Cost (USD)': round(result['words'] * cost_per_word, 2)
		})

	df = pd.DataFrame(data)

	totals = {
		'File Name': 'TOTAL',
		'Type': '',
		'Words': df['Words'].sum(),
		'Cost (USD)': df['Cost (USD)'].sum()
	}
	df = pd.concat([df, pd.DataFrame([totals])], ignore_index=True)
	output_file = f"multi_format_report_{datetime.now().strftime('%Y%m%d-%H%M%S')}.xlsx"

	df.to_excel(output_file, index=False, sheet_name='Word Count')
	print(f"\n✓ Excel report saved: {output_file}")
	return output_file

if __name__ == '__main__':
	print("\n" + "="*70)
	print("MULTI-FORMAT LOCALIZATION WORD COUNTER")
	print("="*70)
	print("Supported formats: JSON, XML, XLF, DOCX, PDF")
	print("="*70)

	folder_path = "/Users/inyoungkim/PycharmProjects/localization-workflow-toolkit/word_counter/test_files"

	results = analyze_folder(folder_path)

	if results:
		display_summary(results, cost_per_word=0.15)
		export_to_excel(results, cost_per_word=0.15)

