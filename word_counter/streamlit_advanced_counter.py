"""
Features:
- Multi-format support (JSON, XML, XLF, DOCX, PDF)
- Language-specific cost calculation (by Reading a rate.json data)
- Excel export
- Configurable rates
"""

import streamlit as st
import os
import json
import xml.etree.ElementTree as ET
from docx import Document
import PyPDF2
import pandas as pd
from datetime import datetime
import io
import json

# CONFIGURATION - Edit rates here
# ============================================================================
LANGUAGE_RATES = {
	'English (EN)': 0.16,
    'French (FR)': 0.16,
    'German (DE)': 0.16,
    'Spanish (ES)': 0.14,
    'Japanese (JA)': 0.20,
    'Simplified Chinese (zh-CN)': 0.11,
}

SOURCE_LANGUAGE = 'Korean (KO)'

# Page Configuration
# ============================================================================

st.set_page_config(
	page_title="Localization Word Counter",
	page_icon="📊",
	layout="wide"
)

# Helper Functions
# ============================================================================

def extract_text_from_json(file):

	def get_strings(obj):
		text = []
		if isinstance(obj, dict):
			for value in obj.values():
				text.extend(get_strings(value))
		elif isinstance(obj, list):
			for item in obj:
				text.extend(get_strings(item))
		elif isinstance(obj, str):
			text.append(obj)
		return text

	try:
		data = json.load(file)
		all_text = get_strings(data)
		return ' '.join(all_text)
	except Exception as e:
		st.error(f"Error reading JSON: {e}")
		return ""


def extract_text_from_xml(file):

	try:
		tree = ET.parse(file)
		root = tree.getroot()
		text = []

		# Check if XLF (only extract source)
		is_xlf = file.name.lower().endswith('.xlf')

		if is_xlf:
			for elem in root.iter():
				if 'source' in elem.tag.lower():
					if elem.text and elem.text.strip():
						text.append(elem.text.strip())
		else:
			for elem in root.iter():
				if elem.text and elem.text.strip():
					text.append(elem.text.strip())

		return ' '.join(text)
	except Exception as e:
		st.error(f"Error reading XML: {e}")
		return ""


def extract_text_from_docx(file):

	try:
		doc = Document(file)
		text = []

		for paragraph in doc.paragraphs:
			if paragraph.text.strip():
				text.append(paragraph.text)

		for table in doc.tables:
			for row in table.rows:
				for cell in row.cells:
					if cell.text.strip():
						text.append(cell.text)

		return ' '.join(text)
	except Exception as e:
		st.error(f"Error reading DOCX: {e}")
		return ""


def extract_text_from_pdf(file):

	try:
		pdf = PyPDF2.PdfReader(file)
		text = []

		for page in pdf.pages:
			page_text = page.extract_text()
			if page_text:
				text.append(page_text)

		return ' '.join(text)
	except Exception as e:
		st.error(f"Error reading PDF: {e}")
		return ""


def count_words_in_file(file):

	filename = file.name
	_, ext = os.path.splitext(filename)
	ext = ext.lower()

	if ext == '.json':
		text = extract_text_from_json(file)
		file_type = 'JSON'
	elif ext in ['.xml', '.xlf']:
		text = extract_text_from_xml(file)
		file_type = 'XML/XLF'
	elif ext == '.docx':
		text = extract_text_from_docx(file)
		file_type = 'DOCX'
	elif ext == '.pdf':
		text = extract_text_from_pdf(file)
		file_type = 'PDF'
	else:
		st.warning(f"⚠️ Unsupported file type: {filename}")
		return None

	words = len(text.split()) if text else 0

	return {
		'filename': filename,
		'file_type': file_type,
		'words': words
	}


def calculate_costs(word_count, selected_languages, rates):

	costs = {}
	total = 0

	for lang in selected_languages:
		rate = rates[lang]
		cost = word_count * rate
		costs[lang] = cost
		total += cost

	return costs, total


def create_excel_report(file_results, selected_languages, language_costs, total_words):
	"""Create Excel report with multiple sheets"""

	# Create Excel writer object in memory
	output = io.BytesIO()

	with pd.ExcelWriter(output, engine='openpyxl') as writer:

		# Sheet 1: File Analysis
		file_data = []
		for result in file_results:
			file_data.append({
				'File Name': result['filename'],
				'Type': result['file_type'],
				'Words': result['words']
			})

		df_files = pd.DataFrame(file_data)

		# Add totals row
		totals_row = {
			'File Name': 'TOTAL',
			'Type': '',
			'Words': df_files['Words'].sum()
		}
		df_files = pd.concat([df_files, pd.DataFrame([totals_row])], ignore_index=True)

		df_files.to_excel(writer, sheet_name='File Analysis', index=False)

		# Sheet 2: Cost by Language
		cost_data = []
		for lang, cost in language_costs.items():
			cost_data.append({
				'Target Language': lang,
				'Rate (USD/word)': LANGUAGE_RATES[lang],
				'Source Words': total_words,
				'Total Cost (USD)': round(cost, 2)
			})

		df_costs = pd.DataFrame(cost_data)

		# Add totals
		cost_totals = {
			'Target Language': 'TOTAL',
			'Rate (USD/word)': '',
			'Source Words': '',
			'Total Cost (USD)': df_costs['Total Cost (USD)'].sum()
		}
		df_costs = pd.concat([df_costs, pd.DataFrame([cost_totals])], ignore_index=True)

		df_costs.to_excel(writer, sheet_name='Cost by Language', index=False)

		# Sheet 3: Summary
		summary_data = {
			'Metric': [
				'Source Language',
				'Total Files',
				'Total Words',
				'Target Languages',
				'Total Cost (USD)',
				'Report Date'
			],
			'Value': [
				SOURCE_LANGUAGE,
				len(file_results),
				total_words,
				', '.join(selected_languages),
				f"${sum(language_costs.values()):,.2f}",
				datetime.now().strftime('%Y-%m-%d %H:%M:%S')
			]
		}
		df_summary = pd.DataFrame(summary_data)
		df_summary.to_excel(writer, sheet_name='Summary', index=False)

		# Format columns
		for sheet_name in writer.sheets:
			worksheet = writer.sheets[sheet_name]
			for column in worksheet.columns:
				max_length = 0
				column_letter = column[0].column_letter
				for cell in column:
					try:
						if len(str(cell.value)) > max_length:
							max_length = len(cell.value)
					except:
						pass
				adjusted_width = min(max_length + 2, 50)
				worksheet.column_dimensions[column_letter].width = adjusted_width

	output.seek(0)
	return output


# Main App
# ============================================================================

st.title("📊 Localization Word Counter")
st.markdown(f"**Source Language:** {SOURCE_LANGUAGE}")
st.markdown("---")

# Sidebar - Configuration
with st.sidebar:
	st.header("⚙️ Configuration")

	st.subheader("Target Languages")
	selected_languages = st.multiselect(
		"Select target languages:",
		options=list(LANGUAGE_RATES.keys()),
		default=list(LANGUAGE_RATES.keys())[:2],  # Default: first 2
		help="Select which languages you want to translate into"
	)

	st.markdown("---")

	st.subheader("💰 Current Rates (USD/word)")
	for lang, rate in LANGUAGE_RATES.items():
		if lang in selected_languages:
			st.markdown(f"**{lang}**: ${rate:.2f}")
		else:
			st.markdown(f"{lang}: ${rate:.2f}")

	st.markdown("---")
	st.caption("💡 Tip: Edit rates in the script's LANGUAGE_RATES dictionary")

# File uploader
st.subheader("📁 Upload Files")
uploaded_files = st.file_uploader(
	"Choose localization files to analyze",
	type=['json', 'xml', 'xlf', 'docx', 'pdf'],
	accept_multiple_files=True,
	help="Supported formats: JSON, XML, XLF, DOCX, PDF"
)

if uploaded_files:

	# Process files
	with st.spinner("Processing files..."):
		results = []

		progress_bar = st.progress(0)
		status_text = st.empty()

		for idx, file in enumerate(uploaded_files):
			status_text.text(f"Processing: {file.name}")
			result = count_words_in_file(file)
			if result:
				results.append(result)
			progress_bar.progress((idx + 1) / len(uploaded_files))

		status_text.empty()
		progress_bar.empty()

	if results:
		# Calculate totals
		total_words = sum(r['words'] for r in results)

		st.success(f"✅ Processed {len(results)} file(s) successfully")

		# Main metrics
		st.markdown("---")
		st.subheader("📈 Overview")

		col1, col2, col3 = st.columns(3)

		with col1:
			st.metric("Total Files", len(results))

		with col2:
			st.metric("Total Words", f"{total_words:,}")

		with col3:
			st.metric("Target Languages", len(selected_languages))

		# File details
		st.markdown("---")
		st.subheader("📄 File Analysis")

		df_files = pd.DataFrame(results)
		df_files = df_files.rename(columns={
			'filename': 'File Name',
			'file_type': 'Type',
			'words': 'Words'
		})

		st.dataframe(
			df_files,
			use_container_width=True,
			hide_index=True
		)

		# Cost calculation
		if selected_languages:
			st.markdown("---")
			st.subheader("💰 Cost Estimation")

			language_costs, total_cost = calculate_costs(
				total_words,
				selected_languages,
				LANGUAGE_RATES
			)

			# Cost breakdown
			cost_data = []
			for lang in selected_languages:
				cost_data.append({
					'Target Language': lang,
					'Rate ($/word)': f"${LANGUAGE_RATES[lang]:.2f}",
					'Words': f"{total_words:,}",
					'Total Cost': f"${language_costs[lang]:,.2f}"
				})

			df_costs = pd.DataFrame(cost_data)

			st.dataframe(
				df_costs,
				use_container_width=True,
				hide_index=True
			)

			# Total cost highlight
			st.markdown("### 💵 Total Translation Cost")
			st.markdown(f"## ${total_cost:,.2f}")

			# Time estimate
			words_per_hour = 250
			total_hours = (total_words * len(selected_languages)) / words_per_hour
			total_days = total_hours / 8

			col1, col2 = st.columns(2)
			with col1:
				st.metric("Estimated Time", f"{total_hours:.1f} hours")
			with col2:
				st.metric("Working Days", f"{total_days:.1f} days")

			# Export to Excel
			st.markdown("---")
			st.subheader("📥 Export Report")

			col1, col2 = st.columns([1, 3])

			with col1:
				if st.button("🔄 Generate Excel Report", type="primary", use_container_width=True):
					with st.spinner("Creating Excel report..."):
						excel_file = create_excel_report(
							results,
							selected_languages,
							language_costs,
							total_words
						)

						filename = f"localization_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

						st.download_button(
							label="📥 Download Excel Report",
							data=excel_file,
							file_name=filename,
							mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
							use_container_width=True
						)

						st.success("✅ Report ready for download!")

			with col2:
				st.info("💡 The Excel report includes:\n- File analysis\n- Cost breakdown by language\n- Summary sheet")

		else:
			st.warning("⚠️ Please select at least one target language to calculate costs.")

	else:
		st.error("❌ No files could be processed successfully.")

else:
	# Welcome message
	st.info("👆 Upload files to get started")

	st.markdown("### 📋 Supported File Formats")
	cols = st.columns(5)
	formats = [
		("📄 JSON", "Localization strings"),
		("📝 XML", "Android resources"),
		("🔤 XLF", "XLIFF translation files"),
		("📰 DOCX", "Word documents"),
		("📕 PDF", "PDF documents")
	]

	for col, (icon, desc) in zip(cols, formats):
		with col:
			st.markdown(f"**{icon}**")
			st.caption(desc)

# Footer
st.markdown("---")
st.caption(f"💡 Edit language rates in the script • Source: {SOURCE_LANGUAGE}")